"""Word 模板套用：识别 {{占位符}} 并用数据 / AI 填充。"""
import io
import json
import re

from docx import Document

from . import llm

_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_placeholders(content):
    """返回模板中所有 {{占位符}} 的名称（去重、排序）。"""
    doc = Document(io.BytesIO(content))
    fields = set()
    for p in _iter_paragraphs(doc):
        for m in _PATTERN.finditer(p.text):
            fields.add(m.group(1).strip())
    return sorted(fields)


def fill_template(content, values):
    """把占位符替换为 values 中的值，返回新的 docx 字节。"""
    doc = Document(io.BytesIO(content))

    def repl(m):
        key = m.group(1).strip()
        return str(values.get(key, m.group(0)))

    for p in _iter_paragraphs(doc):
        if "{{" in p.text:
            p.text = _PATTERN.sub(repl, p.text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_values(text):
    """从 AI 输出中尽力提取 JSON 对象。"""
    try:
        obj = json.loads(text)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        m = re.search(r"\{.*\}", text, re.S)
        if not m:
            return {}
        try:
            return json.loads(m.group(0))
        except Exception:
            return {}


def ai_fill(cfg, content, context="", timeout=600):
    """用大模型为模板占位符生成合适的内容，返回 {变量名: 值}。"""
    fields = extract_placeholders(content)
    if not fields:
        raise RuntimeError("模板中没有找到 {{占位符}}。")
    user = (
        "请为以下模板变量分别生成一句合适的填写内容。\n"
        "模板变量：%s\n"
        "使用背景：%s\n"
        "只输出一个 JSON 对象（key 为变量名，value 为要填进文档的内容），"
        "不要输出任何其他文字。"
        % ("、".join(fields), context or "无")
    )
    text = llm.chat(
        cfg,
        [{"role": "system", "content": "你是办公文档处理助手，只输出指定 JSON。"},
         {"role": "user", "content": user}],
        max_tokens=2048,
        timeout=timeout,
    )
    values = _parse_values(text)
    for f in fields:
        values.setdefault(f, "")
    return values
"""Word 文档处理：AI 生成、摘要、翻译。"""
import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from . import llm
from .md_parse import parse_markdown

WORD_SYSTEM = (
    "你是专业的办公室文档写作助手，擅长中文公文、总结、方案、报告。"
    "你只输出文档正文内容，严格遵守用户给出的输出格式要求，不输出任何解释或额外说明。"
)


def _set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_docx_bytes(blocks, title_text=None):
    """根据解析后的 block 列表生成 docx，返回 bytes。"""
    doc = Document()
    _set_normal_style(doc)
    has_title = False
    for b in blocks:
        t, txt = b.get("type", "para"), b.get("text", "")
        if t in ("title", "h1"):
            doc.add_heading(txt, level=0)
            has_title = True
        elif t == "h2":
            doc.add_heading(txt, level=1)
        elif t == "h3":
            doc.add_heading(txt, level=2)
        elif t == "bullet":
            doc.add_paragraph(txt, style="List Bullet")
        elif t == "numbered":
            doc.add_paragraph(txt, style="List Number")
        else:
            doc.add_paragraph(txt)
    if not has_title and title_text:
        doc.add_heading(title_text, level=0)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_document(cfg, topic, extra=""):
    """根据主题 + 要求生成 Word 文档，返回 (docx_bytes, 原始AI内容)。"""
    user = (
        "请围绕主题【%s】撰写一份结构完整、内容详实的文档。\n"
        "写作要求：%s\n"
        "【输出格式要求】\n"
        "1. 第一行以「# 」开头写文档标题；\n"
        "2. 一级章节用「## 」开头，二级小节用「### 」开头；\n"
        "3. 正文使用普通段落，枚举使用「- 」开头；\n"
        "4. 直接输出文档内容，不要输出任何解释。"
        % (topic, extra or "无")
    )
    content = llm.chat(
        cfg,
        [{"role": "system", "content": WORD_SYSTEM},
         {"role": "user", "content": user}],
        max_tokens=int(cfg.get("max_tokens", 4096)),
        timeout=900,
    )
    blocks = parse_markdown(content)
    return build_docx_bytes(blocks, topic), content


def read_docx_text(content):
    """读取 docx 的字节内容，返回非空段落文本列表。"""
    doc = Document(io.BytesIO(content))
    return [p.text for p in doc.paragraphs if p.text.strip()]


def summarize_document(cfg, content, max_chars=8000):
    """对 Word 文档生成中文摘要。"""
    paras = read_docx_text(content)
    text = "\n".join(paras)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n……（内容过长已截断）"
    user = (
        "请阅读以下文档，生成一份中文摘要，包含：文档主旨、核心要点、结论与建议。"
        "简明扼要，用「-」列表输出。\n\n【文档内容】\n" + text
    )
    return llm.chat(
        cfg,
        [{"role": "system", "content": WORD_SYSTEM},
         {"role": "user", "content": user}],
        max_tokens=2048,
        timeout=600,
    )


def _parse_numbered(text, expect):
    """解析『编号. 内容』格式的翻译结果。"""
    out = []
    lines = [l for l in text.splitlines() if l.strip()]
    for l in lines:
        m = re.match(r"^(\d{1,3})[.、．.)]\s*(.*)$", l.strip())
        if m:
            out.append(m.group(2).strip())
    if not out or len(out) < max(1, len(lines) // 2):
        return lines[:expect]
    return out[:expect]


def translate_document(cfg, content, target="英语"):
    """翻译 Word 文档，返回翻译后 docx 的字节内容。"""
    doc = Document(io.BytesIO(content))
    paras = doc.paragraphs
    texts = [p.text for p in paras]
    idx = [i for i, t in enumerate(texts) if t.strip()]
    blocks = [texts[i] for i in idx]

    translations = []
    for i in range(0, len(blocks), 25):
        sub = blocks[i:i + 25]
        req = "\n".join("%d. %s" % (j + 1, t) for j, t in enumerate(sub))
        user = (
            "请把下面编号的段落逐条翻译成%s。\n"
            "输出格式要求：每条译文占一行，行首为『编号. 』，只输出译文，不输出其他内容。\n\n%s"
            % (target, req)
        )
        out = llm.chat(
            cfg,
            [{"role": "system", "content": "你是专业翻译，只输出译文。"},
             {"role": "user", "content": user}],
            max_tokens=4096,
            timeout=600,
        )
        translations.extend(_parse_numbered(out, len(sub)))

    new_doc = Document()
    _set_normal_style(new_doc)
    ti = 0
    for p in paras:
        if p.text.strip():
            el = new_doc.add_paragraph()
            el.add_run(translations[ti] if ti < len(translations) else p.text)
            if ti < len(translations):
                ti += 1
            try:
                el.style = p.style.name
            except Exception:
                pass
            if p.alignment is not None:
                el.alignment = p.alignment
        else:
            new_doc.add_paragraph()

    buf = io.BytesIO()
    new_doc.save(buf)
    return buf.getvalue()